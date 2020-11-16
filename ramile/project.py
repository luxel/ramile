from docx import Document
from ramile.project_info import ProjectInfo
from ramile.project_processor import ProjectProcessor
from ramile.processors import FileProcessor
import os


class Project(object):
    info = None
    output = True
    files = []
    lines = []

    def __init__(self, project_root, lines_to_extract=3000, output_file='extracted_code.docx', output=True):
        self.info = ProjectInfo(project_root, lines_to_extract)
        self.output = output
        if output:
            self.output_path = self.info.get_output_file_path(output_file)
            # self.output_file = open(
            # self.info.get_output_file_path(output_file), 'w+')
            self.output_file = Document(os.path.join(
                os.path.dirname(__file__), 'data/template.docx'))
            self.paragraph = None
        return

    def run(self, output=True, echo=True):
        if echo:
            print("I'm going to extract %s lines from %s." %
                  (self.info.lines_to_extract, self.info.project_root))
        self.info.lines_extracted = 0
        project_processor = ProjectProcessor(self.info)
        file_processor = FileProcessor()
        # 1. Process and collect the files
        self.files = project_processor.process()
        # 2. Process each file
        for file in self.files:
            for output in file_processor.process(file):
                self.export(output)
                file.extracted_line()
                if self.info.has_extracted_enough_lines():
                    break
            # collect file summary
            self.info.lines_skipped_blank += file.blank_lines
            self.info.lines_skipped_comments += file.comment_lines
            if self.info.has_extracted_enough_lines():
                break
        # self.output_file.close()

        self.write_to_file()

        if echo:
            self.print_summary()

        if not self.info.has_extracted_enough_lines():
            print("Warning!! Not enough source code to extract %s lines!" %
                  self.info.lines_to_extract)
        return

    def print_summary(self):
        print("The extraction is done. Here's the summary:")
        print("Files that contributed to the output:")
        for file in self.files:
            if file.has_extracted_lines():
                print("%s : %s lines" % (file.file_path, file.extracted_lines))
        print("Code was extracted in: %s" % self.output_path)
        print("Total extracted: %s lines" % self.info.lines_extracted)
        print("Wrote to file: %s lines" % len(self.lines))
        print("Total skipped comments: %s lines" %
              self.info.lines_skipped_comments)
        print("Total skipped blank lines: %s lines" %
              self.info.lines_skipped_blank)
        if self.info.lines_extracted > 3000:
            print("Total skipped overflow lines: %s lines" %
                  (self.info.lines_extracted - len(self.lines)))

    def export(self, line):
        max_length_of_line = 60
        appended = 0

        while appended < len(line):
            l = line[appended:appended+max_length_of_line]
            self.lines.append(l)
            self.info.lines_extracted += 1
            appended += len(l)

        return

    def write_to_file(self):
        if self.output:
            if self.paragraph is None:
                self.paragraph = self.output_file.paragraphs[0]

            if self.info.lines_extracted > 3000:
                lines_to_cut = self.info.lines_extracted - 3000
                del self.lines[1501:1501+lines_to_cut]

            i = 0
            for line in self.lines:
                i += 1
                if i < 3000 and not line.endswith('\n'):
                    line += '\n'
                if i == 3000 and line.endswith('\n'):
                    line = line[0:len(line)-1]

                self.paragraph.add_run(line)

            self.output_file.save(self.output_path)
        return
